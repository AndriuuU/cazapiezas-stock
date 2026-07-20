from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
CAPTURES = ROOT / "capturas"
OUTPUT = ROOT / "Manual_usuario_Almacen_Desguace.docx"
LOGO = ROOT.parent.parent / "public" / "logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "20252B"
MUTED = "667085"
AMBER = "F59E0B"
LIGHT_BLUE = "E8EEF5"
LIGHT_AMBER = "FFF6DD"
LIGHT_RED = "FDECEC"
LIGHT_GREEN = "EAF7EF"
BORDER = "D0D5DD"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep(paragraph, with_next=False):
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = with_next


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_run_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True

    if "Tip" not in [style.name for style in doc.styles]:
        tip = doc.styles.add_style("Tip", WD_STYLE_TYPE.PARAGRAPH)
    else:
        tip = doc.styles["Tip"]
    tip.font.name = "Calibri"
    tip.font.size = Pt(10.5)
    tip.paragraph_format.space_after = Pt(0)
    tip.paragraph_format.line_spacing = 1.2


def make_numbering(doc, kind="decimal"):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if kind == "decimal" else "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if kind == "decimal" else "•")
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Symbol")
        fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def restart_num(doc, source_num_id):
    """Create a numbering instance that starts at 1 for one independent list."""
    numbering = doc.part.numbering_part.element
    source_num = next(
        num for num in numbering.findall(qn("w:num"))
        if int(num.get(qn("w:numId"))) == int(source_num_id)
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return new_num_id


def add_list(doc, items, num_id):
    active_num_id = restart_num(doc, num_id)
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, active_num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(label)
            r.bold = True
            p.add_run(detail)
        else:
            p.add_run(item)


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "18")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p = cell.paragraphs[0]
    p.style = doc.styles["Tip"]
    run = p.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    set_keep(p)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_figure(doc, filename, caption, alt_text):
    image_path = CAPTURES / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(image_path), width=Inches(6.35))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep(cap)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths, indent=120)
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(value)
        set_run_font(run, size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=9.5, color=INK)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
    set_table_geometry(table, widths, indent=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_section_heading(doc, title, intro=None, page_break=True):
    p = doc.add_heading(title, level=1)
    p.paragraph_format.page_break_before = page_break
    if intro:
        lead = doc.add_paragraph(intro)
        lead.paragraph_format.space_after = Pt(10)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    decimal_num = make_numbering(doc, "decimal")
    bullet_num = make_numbering(doc, "bullet")

    # Running header/footer.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("CAZAPIEZAS  |  ALMACÉN DESGUACE")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    # Editorial cover.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(22)
    logo_shape = p.add_run().add_picture(str(LOGO), width=Inches(1.25))
    logo_shape._inline.docPr.set("descr", "Logotipo de Cazapiezas Stock")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("GUÍA OPERATIVA")
    set_run_font(kr, size=10, color=AMBER, bold=True)
    kicker.paragraph_format.space_after = Pt(14)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    tr = title.add_run("Manual de usuario")
    set_run_font(tr, size=30, color=DARK_BLUE, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    sr = subtitle.add_run("Almacén Desguace")
    set_run_font(sr, size=18, color=BLUE, bold=True)
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(80)
    dr = desc.add_run("Registro, consulta, fotografías, venta y organización física de piezas")
    set_run_font(dr, size=12, color=MUTED)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Versión 1.0  |  {date.today().strftime('%d/%m/%Y')}")
    set_run_font(mr, size=10, color=MUTED, bold=True)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = audience.add_run("Dirigido al personal de almacén, identificación, fotografía y venta")
    set_run_font(ar, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()

    doc.add_heading("Cómo utilizar este manual", level=1)
    doc.add_paragraph("Este documento explica el recorrido completo de una pieza: desde que entra en el almacén hasta que se publica, reserva, vende, envía o retira. Las capturas son ejemplos reales; el número de piezas y sus datos pueden cambiar.")
    add_callout(doc, "Inicio rápido", "Para registrar una pieza nueva: pulsa Nueva pieza, completa al menos la identificación básica, añade fotografías, elige una ubicación y pulsa Guardar pieza.", LIGHT_GREEN, "237A45")
    doc.add_heading("Contenido", level=2)
    add_list(doc, [
        "Acceso y pantalla principal",
        "Búsqueda, filtros, escáner y paginación",
        "Registro y modificación de piezas",
        "Ficha, fotografías y etiqueta QR",
        "Ubicación física y sugerencias",
        "Acciones rápidas y cambios masivos",
        "Creación y modificación de estanterías por niveles",
        "Flujos recomendados, errores y buenas prácticas",
    ], bullet_num)
    doc.add_heading("Conceptos básicos", level=2)
    add_table(doc, ["Concepto", "Qué significa"], [
        ("Código interno", "Identificador automático con formato CZP-DESGUACE-000000."),
        ("Estado de la pieza", "Condición física o de revisión de la pieza."),
        ("Estado del proceso", "Fase de trabajo: identificar, comprobar, fotografiar, publicar, vender, enviar, etc."),
        ("Ubicación", "Hueco físico con formato DESGUACE-E01-N01-C01."),
        ("Publicado online", "Indica si la pieza se muestra en el canal de venta online."),
    ], [2100, 7260])

    add_section_heading(doc, "1. Acceso y pantalla principal", "La pantalla principal reúne el inventario, los filtros y las acciones de cada pieza.")
    add_figure(doc, "01-listado-principal.png", "Figura 1. Listado principal del Almacén Desguace.", "Listado del almacén con búsqueda, filtros y tabla de piezas.")
    doc.add_heading("Qué información muestra cada fila", level=2)
    add_list(doc, [
        ("Fotos. ", "Miniatura y acceso directo a la galería."),
        ("Referencia. ", "Referencia principal y nombre de la pieza."),
        ("Coche. ", "Datos del vehículo compatible al desplegar Ver coche."),
        ("Precio y fecha. ", "Precio de venta y fecha de entrada."),
        ("Ubicación. ", "Hueco exacto o aviso Sin ubicar."),
        ("Estados. ", "Condición de la pieza y fase del proceso."),
        ("Online. ", "Confirma si la pieza está publicada online."),
        ("Acciones. ", "Abre todas las operaciones disponibles para esa pieza."),
    ], bullet_num)
    add_callout(doc, "Consejo", "Antes de crear una pieza, busca por referencia principal u OEM para evitar duplicados.", LIGHT_AMBER, "9A6700")

    add_section_heading(doc, "2. Buscar, filtrar y localizar piezas")
    doc.add_heading("Búsqueda libre", level=2)
    doc.add_paragraph("Escribe código interno, nombre, referencia, marca, modelo o motor. Puedes usar varias palabras; todas deberán aparecer en alguno de los campos buscables.")
    doc.add_heading("Filtros disponibles", level=2)
    add_list(doc, [
        "Categoría de la pieza.",
        "Estado físico o de revisión.",
        "Estado del proceso.",
        "Publicadas online o no publicadas.",
        "Ubicación completa, estantería o fragmento, por ejemplo E01.",
        "Orden por fecha, nombre, referencia, ubicación o precio.",
    ], bullet_num)
    doc.add_heading("Paginación", level=2)
    doc.add_paragraph("Puedes mostrar 10, 25, 50 o 100 filas por página. Utiliza los controles inferiores para avanzar, retroceder o saltar directamente a una página.")
    doc.add_heading("Buscar con cámara o lector", level=2)
    add_list(doc, [
        "Pulsa Escanear.",
        "En un móvil compatible, pulsa Activar cámara y centra el código en el recuadro.",
        "También puedes usar un lector USB/Bluetooth o escribir la referencia manualmente.",
        "Al detectar el código, el sistema lo coloca en la búsqueda del inventario.",
    ], decimal_num)
    add_figure(doc, "09-escanear-referencia.png", "Figura 2. Búsqueda mediante cámara, lector o escritura manual.", "Cuadro para escanear una referencia con cámara o lector físico.")
    add_callout(doc, "Si la cámara no funciona", "Utiliza un lector USB/Bluetooth o escribe la referencia. No es necesario conceder permiso de cámara para esas alternativas.", LIGHT_AMBER, "9A6700")

    add_section_heading(doc, "3. Registrar una pieza nueva", "El formulario admite borradores incompletos y comprueba automáticamente los requisitos cuando se intenta publicar.")
    add_figure(doc, "04-nueva-pieza.png", "Figura 3. Formulario de alta de una pieza.", "Formulario para crear una pieza con datos, identificación y vehículo compatible.")
    doc.add_heading("Paso 1. Datos e identificación", level=2)
    add_list(doc, [
        ("Nombre de la pieza. ", "Nombre reconocible, por ejemplo FARO DERECHO."),
        ("Descripción. ", "Estado, lado, color, desperfectos y observaciones útiles."),
        ("Categoría. ", "Familia general que facilita filtros y reglas de estantería."),
        ("Marca de la pieza. ", "Fabricante del componente, si se conoce."),
        ("Referencias. ", "Principal, OEM y equivalentes. Es recomendable completar al menos una referencia."),
        ("Vehículo. ", "Marca, modelo, motorización, código motor y años compatibles."),
        ("Procedencia. ", "Origen de la pieza o información de trazabilidad."),
    ], bullet_num)
    doc.add_heading("Paso 2. Estado, cantidad y precios", level=2)
    add_list(doc, [
        "Elige el estado real de la pieza.",
        "Selecciona la fase de trabajo en Estado del proceso.",
        "Indica la cantidad disponible y los precios sin valores negativos.",
        "Comprueba que Año hasta no sea anterior a Año desde.",
        "Marca Publicada online solo cuando la ficha esté preparada para venta.",
    ], decimal_num)
    doc.add_heading("Paso 3. Ubicación", level=2)
    doc.add_paragraph("Puedes aceptar la recomendación, elegir un hueco libre, escribir una ubicación exacta o dejarla vacía para ubicar la pieza después.")
    add_callout(doc, "Formato", "La ubicación completa se forma así: DESGUACE-E01-N03-C05 = estantería E01, nivel 03, hueco 05.", LIGHT_BLUE, BLUE)
    doc.add_heading("Paso 4. Fotografías", level=2)
    doc.add_paragraph("Pulsa Añadir fotos y selecciona una o varias imágenes. Solo se admiten archivos de imagen de hasta 10 MB cada uno. Puedes quitar una previsualización antes de guardar.")
    doc.add_heading("Guardar como borrador o publicar", level=2)
    add_table(doc, ["Opción", "Comportamiento"], [
        ("Borrador", "Puedes guardar campos incompletos con un estado de proceso anterior a Lista para publicar."),
        ("Lista para publicar / Publicada", "Se exigen nombre, referencia, estado, precio de venta, ubicación, cantidad y al menos una fotografía."),
    ], [2700, 6660])
    add_callout(doc, "Importante", "Si falta algún requisito para publicar, el sistema muestra un aviso con la lista exacta de datos pendientes. Completa esos datos y vuelve a guardar.", LIGHT_RED, "9B1C1C")

    add_section_heading(doc, "4. Consultar y modificar una ficha")
    doc.add_heading("Contenido de la ficha", level=2)
    add_list(doc, [
        "Datos básicos, categoría, marca, procedencia y fecha de entrada.",
        "Referencias principal, OEM y equivalentes.",
        "Vehículo compatible, motorización, código motor y años.",
        "Estado, cantidad, precios, proceso y publicación online.",
        "Ubicación física exacta.",
        "Galería completa de fotografías.",
    ], bullet_num)
    doc.add_heading("Editar", level=2)
    doc.add_paragraph("Pulsa Editar en la ficha o en el panel de acciones. El formulario es el mismo que el alta, pero aparece cargado con los datos existentes. Revisa y pulsa Guardar pieza.")
    doc.add_heading("Imprimir etiqueta QR", level=2)
    add_list(doc, [
        "Pulsa Imprimir etiqueta QR.",
        "El navegador abre la vista de impresión con código interno, nombre, referencia, ubicación, estado y QR.",
        "El QR lleva directamente a la ficha de la pieza.",
        "Si no se abre la impresión, permite las ventanas emergentes para esta aplicación.",
    ], decimal_num)

    add_section_heading(doc, "5. Gestionar fotografías")
    doc.add_paragraph("La fotografía principal es la que se utiliza como imagen destacada. En la galería se reconoce por la estrella y el borde resaltado.")
    add_list(doc, [
        ("Subir fotos. ", "Permite añadir varias imágenes desde la ficha o mientras se crea la pieza."),
        ("Hacer principal. ", "Pulsa la estrella de la imagen que deseas destacar."),
        ("Copiar enlace público. ", "Copia la dirección pública de la fotografía para usarla en otros canales."),
        ("Eliminar. ", "Pulsa la papelera y confirma. La eliminación es definitiva."),
        ("Ver fotos. ", "Desde el listado abre una galería ampliada sin abandonar el inventario."),
    ], bullet_num)
    add_callout(doc, "Buena práctica", "Fotografía la referencia, el conector, la vista general y cualquier defecto. Evita imágenes oscuras o desenfocadas.", LIGHT_GREEN, "237A45")

    add_section_heading(doc, "6. Ubicar físicamente una pieza", "El sistema cruza el nombre, la descripción y la categoría con las reglas de las estanterías.")
    add_figure(doc, "03-ubicar-pieza.png", "Figura 5. Propuesta automática de ubicación.", "Ventana de colocación con ubicación recomendada y alternativas.")
    doc.add_heading("Aceptar la recomendación", level=2)
    add_list(doc, [
        "Abre Acciones y pulsa Ubicar.",
        "Comprueba la estantería, el nivel y el hueco propuestos.",
        "Coloca físicamente la pieza en ese hueco.",
        "Pulsa Sí, colocada ahí para guardar la ubicación real.",
    ], decimal_num)
    doc.add_heading("Usar otra ubicación", level=2)
    doc.add_paragraph("Pulsa No, está en otro sitio. Puedes elegir una alternativa o escribir la ubicación completa. El sistema comprueba que la estantería exista, esté activa, tenga espacio y que el hueco no esté ocupado.")
    doc.add_heading("Registrar que no se pudo colocar", level=2)
    doc.add_paragraph("Pulsa No pude colocarla, explica el motivo y registra la incidencia. Esta opción sirve para estanterías llenas, piezas demasiado grandes o accesos bloqueados.")
    add_callout(doc, "Recalcular", "Si alguien ha ocupado un hueco mientras la ventana estaba abierta, pulsa Recalcular para obtener una propuesta actualizada.", LIGHT_BLUE, BLUE)

    add_section_heading(doc, "7. Acciones rápidas y cambios masivos", page_break=False)
    add_figure(doc, "02-acciones-pieza.png", "Figura 6. Panel de acciones disponibles para una pieza.", "Fila ampliada con acciones de consulta, edición, ubicación, publicación y venta.")
    doc.add_heading("Acciones de una pieza", level=2)
    add_table(doc, ["Acción", "Resultado"], [
        ("Ver ficha", "Abre toda la información y la galería."),
        ("Editar", "Permite corregir datos, estados, precios y ubicación."),
        ("Ubicar", "Calcula una ubicación recomendada y registra el resultado."),
        ("Ver fotos / Subir fotos", "Consulta la galería o añade nuevas imágenes."),
        ("Publicar", "Cambia la pieza a Publicada si cumple los requisitos."),
        ("Reservar", "Marca la pieza como Reservada."),
        ("Vendida", "Marca la pieza como Vendida; requiere confirmación."),
        ("Enviada", "Marca el proceso como Enviada."),
        ("Retirar", "Retira la pieza del flujo de venta; requiere confirmación."),
    ], [2000, 7360])
    doc.add_heading("Modificar varias piezas", level=2)
    add_figure(doc, "08-cambio-masivo.png", "Figura 7. Barra de modificación masiva tras seleccionar una pieza.", "Listado con una pieza seleccionada y controles de cambio masivo.")
    add_list(doc, [
        "Marca las casillas de las piezas que deseas modificar o selecciona toda la página.",
        "Elige el dato: estado de la pieza, estado del proceso o publicación online.",
        "Selecciona el nuevo valor.",
        "Pulsa Aplicar y confirma el número de piezas afectadas.",
    ], decimal_num)
    add_callout(doc, "Precaución", "Los cambios masivos se aplican a todas las piezas seleccionadas, aunque estén en páginas diferentes. Revisa el contador antes de confirmar.", LIGHT_RED, "9B1C1C")

    add_section_heading(doc, "8. Organizar estanterías")
    add_figure(doc, "06-estanterias.png", "Figura 8. Resumen de estanterías, ocupación y disponibilidad.", "Tarjetas de estanterías con capacidad, huecos libres y controles de estado.")
    doc.add_heading("Información de cada estantería", level=2)
    add_list(doc, [
        "Código y nombre descriptivo.",
        "Número de niveles y huecos por nivel.",
        "Capacidad máxima, ocupados, libres y porcentaje de ocupación.",
        "Siguiente hueco disponible.",
        "Reglas generales o grupos de niveles configurados.",
        "Estado activa, inactiva, disponible o llena.",
    ], bullet_num)
    doc.add_heading("Controles rápidos", level=2)
    add_list(doc, [
        ("Editar. ", "Abre el formulario completo sin perder las piezas ya ubicadas."),
        ("Marcar llena. ", "Bloquea temporalmente nuevas sugerencias aunque queden huecos."),
        ("Marcar con espacio. ", "Quita el bloqueo manual."),
        ("Desactivar. ", "Evita que la estantería participe en las sugerencias."),
        ("Activar. ", "Vuelve a incluirla en las sugerencias."),
    ], bullet_num)

    add_section_heading(doc, "9. Crear o modificar una estantería por niveles", "Una misma estantería puede reservar grupos de niveles para tipos de pieza diferentes.")
    add_figure(doc, "07-formulario-estanteria.png", "Figura 9. Formulario de estantería dividido en tres pasos.", "Formulario para identificar una estantería, definir tamaño y repartir piezas por niveles.")
    doc.add_heading("Paso 1. Identificar", level=2)
    add_list(doc, [
        "Código único con formato E01, E02, E03...",
        "Nombre fácil de reconocer.",
        "Descripción opcional con pasillo, zona u observaciones.",
    ], decimal_num)
    doc.add_heading("Paso 2. Definir el tamaño", level=2)
    add_list(doc, [
        "Número de niveles físicos.",
        "Número de huecos disponibles en cada nivel.",
        "Capacidad que se utilizará. No puede superar los huecos físicos calculados.",
    ], decimal_num)
    doc.add_heading("Paso 3. Repartir piezas por niveles", level=2)
    doc.add_paragraph("Crea un grupo para cada tipo de pieza. Indica Desde nivel, Hasta nivel y Qué piezas van aquí. En Palabras adicionales puedes añadir sinónimos o términos frecuentes.")
    add_callout(doc, "Ejemplo solicitado", "Estantería E01 con 4 niveles: Grupo 1 = niveles 1 a 2, contenido Faros; Grupo 2 = niveles 3 a 4, contenido Retrovisores.", LIGHT_GREEN, "237A45")
    add_table(doc, ["Grupo", "Desde", "Hasta", "Contenido", "Palabras adicionales"], [
        ("1", "1", "2", "Faros", "faro, piloto, óptica, antiniebla"),
        ("2", "3", "4", "Retrovisores", "retrovisor, espejo, cristal espejo"),
    ], [900, 900, 900, 2400, 4260], header_fill=LIGHT_AMBER)
    add_list(doc, [
        "Pulsa Añadir otro grupo de niveles para crear más divisiones.",
        "Pulsa Quitar para eliminar un grupo que ya no necesites.",
        "Un mismo nivel no puede aparecer en dos grupos.",
        "Los niveles indicados deben estar dentro del número total de niveles.",
        "Guarda con Crear estantería o Guardar cambios.",
    ], bullet_num)
    add_callout(doc, "Cómo decide el sistema", "Busca coincidencias entre el nombre, descripción y categoría de la pieza y el contenido o palabras adicionales del grupo. Después elige el primer hueco libre dentro de esos niveles.", LIGHT_BLUE, BLUE)
    doc.add_heading("Errores del formulario", level=2)
    doc.add_paragraph("Si existe un código duplicado, una capacidad incorrecta, niveles solapados o un intervalo fuera de rango, el aviso aparece dentro del formulario y permanece visible. Corrige el dato indicado y vuelve a guardar.")

    add_section_heading(doc, "10. Flujos de trabajo recomendados")
    doc.add_heading("Entrada rápida como borrador", level=2)
    add_list(doc, [
        "Crea la pieza con nombre provisional, procedencia y referencia si está disponible.",
        "Déjala en Pendiente de identificar o Pendiente de comprobar.",
        "Guarda y completa después la ficha, las fotos y la ubicación.",
    ], decimal_num)
    doc.add_heading("Preparar una pieza para venta", level=2)
    add_list(doc, [
        "Confirma identificación, referencia, estado y compatibilidad.",
        "Añade precio de venta, cantidad y fotografías claras.",
        "Ubica físicamente la pieza y confirma el hueco.",
        "Selecciona Lista para publicar para revisar requisitos.",
        "Cuando esté lista, utiliza Publicar y comprueba el indicador Online.",
    ], decimal_num)
    doc.add_heading("Venta y expedición", level=2)
    add_list(doc, [
        "Reserva la pieza para evitar que se gestione como disponible.",
        "Cuando se complete la venta, pulsa Vendida y confirma.",
        "Al preparar la salida, utiliza el estado Pendiente de envío.",
        "Después de expedirla, pulsa Enviada.",
    ], decimal_num)
    doc.add_heading("Reorganizar el almacén", level=2)
    add_list(doc, [
        "Crea o edita las reglas de estantería antes de ubicar un lote grande.",
        "Marca como llena una zona temporalmente bloqueada.",
        "Recalcula la ubicación de cada pieza antes de colocarla.",
        "Usa el filtro de ubicación para revisar el contenido de una estantería.",
    ], decimal_num)

    add_section_heading(doc, "11. Solución de problemas y buenas prácticas")
    add_table(doc, ["Situación", "Qué hacer"], [
        ("No puedo publicar", "Completa nombre, referencia, estado, precio, ubicación, cantidad y al menos una fotografía."),
        ("No hay ubicación recomendada", "Revisa reglas, palabras, estanterías activas, capacidad y bloqueos manuales."),
        ("El hueco está ocupado", "Pulsa Recalcular o selecciona otra ubicación libre."),
        ("La cámara no funciona", "Usa lector USB/Bluetooth o escritura manual."),
        ("No se abre la etiqueta", "Permite ventanas emergentes y vuelve a pulsar Imprimir etiqueta QR."),
        ("No puedo copiar la foto", "Comprueba que el almacenamiento de fotografías esté configurado como público."),
        ("La fotografía no sube", "Verifica que sea una imagen y que no supere 10 MB."),
        ("Error al guardar estantería", "Lee el aviso visible: revisa código, capacidad, intervalos y niveles repetidos."),
    ], [2800, 6560])
    doc.add_heading("Buenas prácticas diarias", level=2)
    add_list(doc, [
        "No marques una pieza como colocada hasta haberla dejado físicamente en el hueco indicado.",
        "Usa nombres claros y añade sinónimos a las reglas de niveles.",
        "Mantén una sola fotografía principal y elimina imágenes repetidas o incorrectas.",
        "Utiliza referencias OEM siempre que estén disponibles.",
        "Revisa la cantidad y el estado después de una venta o devolución.",
        "Evita reutilizar una ubicación sin comprobar antes que esté libre.",
    ], bullet_num)

    add_section_heading(doc, "12. Referencia rápida")
    add_table(doc, ["Necesito...", "Ruta rápida"], [
        ("Crear una pieza", "Almacén Desguace > Nueva pieza"),
        ("Buscar una referencia", "Buscar y filtrar > escribir o Escanear"),
        ("Modificar una pieza", "Acciones > Editar"),
        ("Ver o añadir fotos", "Acciones > Ver fotos / Subir fotos"),
        ("Imprimir etiqueta", "Ver ficha > Imprimir etiqueta QR"),
        ("Colocar una pieza", "Acciones > Ubicar"),
        ("Cambiar varias piezas", "Seleccionar casillas > elegir dato y valor > Aplicar"),
        ("Crear estantería", "Organizar estanterías > Nueva estantería"),
        ("Separar por niveles", "Editar estantería > Reparte las piezas por niveles"),
        ("Bloquear una estantería", "Organizar estanterías > Marcar llena o Desactivar"),
    ], [3000, 6360])
    add_callout(doc, "Regla de oro", "La ubicación registrada debe coincidir siempre con el lugar físico real de la pieza.", LIGHT_AMBER, "9A6700")
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(24)
    run = closing.add_run("Fin del manual · Cazapiezas - Almacén Desguace")
    set_run_font(run, size=10, color=MUTED, italic=True)

    # Core properties and accessibility metadata.
    doc.core_properties.title = "Manual de usuario - Almacén Desguace"
    doc.core_properties.subject = "Uso completo del módulo Almacén Desguace"
    doc.core_properties.author = "Cazapiezas"
    doc.core_properties.keywords = "almacén, desguace, piezas, estanterías, manual"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
